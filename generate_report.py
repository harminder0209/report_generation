#!/usr/bin/env python3
"""
generate_report.py
==================
Generate a BRD-style "Per-Report Optimization Analysis" section from a Power BI
semantic-model JSON (TMDL / .bim style, as exported from the Power BI Service).

The output mirrors section 3.x of "BRD - Power BI Optimization.docx":

    3.x  <Report Name>
      - Table Inventory          (tables, type, columns)
      - DAX Measure Inventory
      - Relationship Map         (with many-to-many detection)
      - Parameters
      Transformation Inventory   (parsed from each partition's M / native SQL)
      Performance Analysis       (rule-based issue detection)
      Performance Mitigations    (issue -> Gold-rebuild mitigation)

Everything is *derived from the model JSON* — no hand entry. Usage stats
(views/users) and ownership are not in the model and are left as placeholders.

Usage:
    .venv/bin/python generate_report.py database.json
    .venv/bin/python generate_report.py database.json -o report.docx
    .venv/bin/python generate_report.py database.json --format md -o report.md
"""

import argparse
import json
import re
import sys
from pathlib import Path

# --------------------------------------------------------------------------- #
#  Colours used for cell shading (matches the BRD's red / amber / green key)  #
# --------------------------------------------------------------------------- #
RED = "F4CCCC"      # critical blocker
AMBER = "FCE5CD"    # significant inefficiency
GREEN = "D9EAD3"    # clean
HEADER = "D9D9D9"   # table header row

AUTO_TABLE_RE = re.compile(r"^(LocalDateTable_|DateTableTemplate_)", re.I)


# --------------------------------------------------------------------------- #
#  Model loading / light abstraction                                          #
# --------------------------------------------------------------------------- #
def load_model(path):
    with open(path, encoding="utf-8") as fh:
        doc = json.load(fh)
    return doc.get("model", doc), doc.get("name", Path(path).stem)


def partition_expression(table):
    """Return the M/SQL expression text of a table's first partition."""
    parts = table.get("partitions") or []
    if not parts:
        return ""
    src = parts[0].get("source", {}) or {}
    expr = src.get("expression", "")
    if isinstance(expr, list):
        expr = "\n".join(expr)
    return expr or ""


def user_tables(model):
    """All tables except Power BI's auto-generated date tables."""
    return [t for t in model.get("tables", []) if not AUTO_TABLE_RE.match(t["name"])]


def safe_filename(name):
    """Turn a report name into a filesystem-safe file stem."""
    stem = re.sub(r'[<>:"/\\|?*]+', "_", name).strip().rstrip(". ")
    stem = re.sub(r"\s+", " ", stem)
    return stem or "report"


# --------------------------------------------------------------------------- #
#  Derivations                                                                #
# --------------------------------------------------------------------------- #
def column_label(col):
    """Column display name, suffixed with '(Agg)' when it has a default agg."""
    name = col["name"]
    if (col.get("summarizeBy") or "none").lower() not in ("none", "", None):
        return f"{name} (Agg)"
    return name


def classify_table(table, model):
    """
    Heuristic table-role classification, derived from name + SQL + how the
    table participates in relationships. Mirrors the BRD's Type column.
    """
    name = table["name"]
    sql = partition_expression(table)
    low = name.lower()

    rels = model.get("relationships", [])
    out_many = any(
        r.get("fromTable") == name and (r.get("toCardinality") == "many")
        for r in rels
    )

    if name in ("RefreshDate",) or "pbi_settings" in sql.lower():
        return "Metadata"
    if "recordtypefr" in low or (low.endswith("fr_asset") and "rel" not in low):
        return "Lookup"
    if "bridge" in low or "UNPIVOT" in sql.upper():
        return "Bridge"
    if re.search(r"(doc\w*rel|rel_asset)$", low) and "ate" not in low:
        return "Relationship"
    # Fact = document-grain tables
    if low.startswith("docs_") or low.startswith("docsprop"):
        return "Fact"
    # Dimensions: hierarchies and the ATE_* attribute tables
    if "hierarchy" in low or low.startswith("ate_"):
        return "Dimension"
    if out_many:
        return "Bridge"
    return "Dimension"


def detect_duplicate_facts(model):
    """Find fact tables that look like duplicates (same grain, FR variant)."""
    facts = [t["name"] for t in user_tables(model) if classify_table(t, model) == "Fact"]
    dups = []
    for a in facts:
        for b in facts:
            if a < b and a.split("_")[0].rstrip("Prop") in b or (
                "Docs" in a and "Docs" in b and a != b
            ):
                dups.append((a, b))
    return list({tuple(sorted(d)) for d in dups})


def relationship_rows(model):
    """
    Build the Relationship Map rows. Cardinality + colour derived from the
    crossFiltering + toCardinality fields in the model.
    """
    rows = []
    for r in model.get("relationships", []):
        ft, fc = r.get("fromTable"), r.get("fromColumn")
        tt = r.get("toTable")
        if AUTO_TABLE_RE.match(tt or "") or AUTO_TABLE_RE.match(ft or ""):
            continue
        to_card = r.get("toCardinality")
        from_card = r.get("fromCardinality")
        xfilter = r.get("crossFilteringBehavior")

        if to_card == "many":
            card, issue, colour = (
                "Many : Many",
                "M:M — cross-filter expansion resolved at query time; expensive on every visual interaction.",
                RED,
            )
        elif from_card == "one":
            card, issue, colour = (
                "1 : 1",
                "1:1 — candidate to fold into the related table as a column (avoid extra join).",
                AMBER,
            )
        else:
            card, issue, colour = ("1 : Many", "Clean — standard star join.", GREEN)

        if xfilter == "bothDirections" and colour == GREEN:
            issue = "Bi-directional filter — review; can cause ambiguity / extra cost."
            colour = AMBER

        rows.append((f"{ft} → {tt}", card, issue, colour))
    return rows


# --- Transformation inventory ---------------------------------------------- #
def analyse_sql(table):
    """
    Inspect a partition's M / native SQL and return (logic_type, detail, flags).
    flags is a set of issue keys used later by the performance analysis.
    """
    sql = partition_expression(table)
    up = sql.upper()
    flags = set()
    details = []
    logic = []

    if not sql:
        return "Unknown", "No partition expression found.", flags

    native = "Sql.Database" in sql
    is_m_shaped = any(
        k in sql for k in ("Table.NestedJoin", "Table.AddColumn", "Table.ExpandTableColumn")
    )

    # temp-table recursive hierarchy
    temps = sorted(set(re.findall(r"#t\d+", sql)))
    if temps:
        logic.append("Recursive CTE")
        details.append(
            f"Hierarchy flattened using {len(temps)} temp tables ({temps[0]}–{temps[-1]}) "
            "with iterative JOINs; CREATE/INSERT/DROP per refresh."
        )
        flags.add("recursive_temp")

    if "UNPIVOT" in up:
        logic.append("UNPIVOT")
        details.append("UNPIVOTs id columns into a single key — creates a many-to-many bridge.")
        flags.add("unpivot_mm")

    # 3-way union existence filter
    if up.count("EXISTS") >= 1 and "UNION" in up:
        logic.append("Existence filter (UNION)")
        n_union = up.count("UNION ALL") + up.count("UNION SELECT") + up.count("\nUNION")
        details.append(
            f"WHERE EXISTS over a UNION of subqueries — correlated scan of relationships "
            f"{max(2, up.count('EXISTS'))}× per refresh."
        )
        flags.add("union_subquery")
    elif "EXISTS" in up:
        logic.append("EXISTS subquery")
        details.append("EXISTS check confirms linkage via the asset hierarchy.")

    if "PIVOT" in up and "UNPIVOT" not in up:
        logic.append("PIVOT")
        details.append("PIVOT + COALESCE builds a display label.")

    if "COALESCE" in up and "PIVOT" not in up:
        logic.append("COALESCE")

    # translated / bilingual table-valued functions
    if "_translated" in sql.lower() or re.search(r"-\s*FR", sql) or "'fr'" in sql.lower():
        details.append("Calls translated (FR) views/functions — bilingual columns.")
        flags.add("bilingual_fr")

    # nested joins done in M (not server-side)
    if is_m_shaped:
        logic.append("M merge")
        details.append("Power Query NestedJoin / ExpandTableColumn — shaping done client-side.")
        flags.add("m_merge")

    # CommandTimeout work-arounds
    m = re.search(r"CommandTimeout\s*=\s*#duration\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", sql)
    if m:
        d, h, mi, s = map(int, m.groups())
        minutes = d * 1440 + h * 60 + mi + s / 60
        details.append(f"⚠ CommandTimeout = {int(minutes)} min — masks a slow query.")
        flags.add("command_timeout")

    # class codes referenced (useful colour for the detail text)
    codes = sorted(set(re.findall(r"\bARE\d{3}\b", sql)))
    if codes:
        details.append(f"Relationship class codes: {', '.join(codes)}.")

    if native and not logic:
        logic.append("Filtered JOIN")
    if not native and not logic:
        logic.append("Power Query (M)")

    source = "SQL (native)" if native and not is_m_shaped else (
        "SQL + M merge" if native and is_m_shaped else "Power Query"
    )
    logic_type = " + ".join(dict.fromkeys(logic)) or "Passthrough"
    detail = " ".join(details) or "Native SQL passed via Sql.Database()."
    return source, logic_type, detail, flags


def transformation_rows(model):
    rows = []
    all_flags = {}
    for t in sorted(user_tables(model), key=lambda x: x["name"].lower()):
        source, logic, detail, flags = analyse_sql(t)
        rows.append((t["name"], source, logic, detail))
        all_flags[t["name"]] = flags
    return rows, all_flags


# --- Performance analysis + mitigations ------------------------------------ #
def measures(model):
    out = []
    for t in model.get("tables", []):
        for me in t.get("measures", []):
            expr = me.get("expression", "")
            if isinstance(expr, list):
                expr = " ".join(expr)
            out.append((me["name"], t["name"], " ".join(expr.split())))
    return out


def parameters(model):
    out = []
    for e in model.get("expressions", []):
        kind = (e.get("kind") or "").lower()
        expr = e.get("expression", "")
        if isinstance(expr, list):
            expr = "\n".join(expr)
        is_param = e.get("kind") == "m" and "IsParameterQuery" in expr or kind == "m"
        # take the first quoted literal as "current value"
        m = re.search(r'"([^"]+)"', expr)
        val = m.group(1) if m else expr.strip()[:80]
        flag = ""
        if "dev" in val.lower():
            flag = "  (⚠ flagged: dev server / environment)"
        out.append((e["name"], val + flag))
    return out


# Catalogue of issues; each entry = (key, title, root-cause template, impact, priority, colour)
ISSUE_LIBRARY = {
    "recursive_temp": (
        "Recursive hierarchy built in SQL via temp tables",
        "{tables} run a multi-iteration loop with CREATE/INSERT/DROP per refresh. "
        "No incremental capability — full reload every time.",
        "High CU, long refresh time", "Critical", RED,
    ),
    "unpivot_mm": (
        "Many-to-many relationships in the model",
        "{tables} UNPIVOT creates M:M relationships. Power BI resolves M:M at query "
        "time via cross-filter expansion — expensive on every visual interaction.",
        "Slow visual load, high CU per interaction", "Critical", RED,
    ),
    "union_subquery": (
        "Correlated subquery with multi-way UNION",
        "{tables} use WHERE EXISTS over a UNION of subqueries, forcing repeated full "
        "scans of the relationships table per refresh.",
        "Long refresh, high server I/O", "High", AMBER,
    ),
    "dup_fact": (
        "Duplicated document fact tables",
        "{tables} query overlapping document sets and are joined/merged in Power Query. "
        "The same data is loaded more than once.",
        "Double memory footprint, double refresh time", "High", AMBER,
    ),
    "bilingual_fr": (
        "Bilingual (FR) columns duplicated across tables",
        "Separate '*- FR' columns and translated table-valued functions appear in "
        "{tables}. Each is an extra SQL call to a translated view/function.",
        "Extra joins, memory overhead, maintenance risk", "Medium", AMBER,
    ),
    "no_incremental": (
        "No incremental refresh — full reload every run",
        "No RangeStart/RangeEnd parameters detected. Every refresh reloads all rows "
        "from all tables regardless of what changed.",
        "CU waste on unchanged data", "High", AMBER,
    ),
    "command_timeout": (
        "Long CommandTimeout values on multiple queries",
        "{tables} set a long CommandTimeout (#duration). This masks slow queries "
        "instead of fixing them.",
        "Masks performance problems, extends refresh window", "Medium", AMBER,
    ),
    "m_merge": (
        "Client-side joins done in Power Query (M)",
        "{tables} use NestedJoin/ExpandTableColumn in M rather than pushing the join "
        "to the server, preventing query folding.",
        "Slower refresh, no fold to source", "Medium", AMBER,
    ),
    "dev_server": (
        "Dev server / environment referenced",
        "A parameter points to a dev source. Dev environments are not sized for "
        "production query loads.",
        "Unreliable performance baselines", "Medium", AMBER,
    ),
}

# Mitigation templates keyed by issue
MITIGATION_LIBRARY = {
    "recursive_temp": (
        "Replace recursive SQL hierarchy with DAX PATH/PATHITEM in the Gold model",
        "1. Store parent_object_id in Gold as a simple column. 2. Add a DAX column "
        "Path = PATH(id, parent). 3. Derive levels via PATHITEM. 4. Materialise in Gold ETL.",
        "Eliminates temp tables; levels computed once at load. PATH is optimised for VertiPaq.",
    ),
    "unpivot_mm": (
        "Resolve M:M with a proper bridge table in Gold",
        "1. Pre-join the dimensions into a single resolved bridge fact in Gold ETL. "
        "2. Model standard 1:M from each dimension to the bridge. 3. Remove UNPIVOT from M.",
        "Eliminates M:M. 1:M joins resolve in microseconds vs seconds for cross-filter expansion. Biggest single gain.",
    ),
    "union_subquery": (
        "Pre-materialise eligibility in Gold",
        "1. Create a Gold view pre-computing linked rows across all paths. 2. Replace the "
        "UNION correlated subquery with a simple SELECT. 3. Index the class_code columns.",
        "Converts a correlated subquery (N full scans) into a single indexed lookup.",
    ),
    "dup_fact": (
        "Consolidate duplicated facts into one Gold fact table",
        "1. Combine the overlapping facts into a single Gold table with EN + FR columns. "
        "2. Remove the NestedJoin step. 3. Retire the separate FR lookup table.",
        "Halves the document-data memory footprint; removes a SQL call and a merge per refresh.",
    ),
    "bilingual_fr": (
        "Model bilingual columns with a language attribute, not duplicate columns",
        "1. Store EN/FR as paired columns on one row, or a language dimension. "
        "2. Remove separate *_FR tables and translated function calls.",
        "Eliminates per-refresh calls to translated table-valued functions; simplifies lineage.",
    ),
    "no_incremental": (
        "Enable incremental refresh with RangeStart/RangeEnd",
        "1. Add RangeStart/RangeEnd parameters. 2. Filter on a watermark timestamp. "
        "3. Configure a rolling partition policy. 4. Requires a reliable Gold watermark column.",
        "Only changed partitions refresh; CU drops proportionally to the change rate.",
    ),
    "command_timeout": (
        "Remove long CommandTimeout work-arounds",
        "1. Fix the underlying slow queries in Gold (indexes, pre-materialised views). "
        "2. Remove the overrides once queries run fast. 3. Set a standard governance timeout.",
        "Long timeouts let runaway queries hold connections and block other refreshes.",
    ),
    "m_merge": (
        "Push joins to the server / Gold instead of Power Query",
        "1. Move NestedJoin logic into a Gold SQL view. 2. Import the pre-joined result. "
        "3. Keep M steps foldable.",
        "Restores query folding; reduces client-side refresh work and memory.",
    ),
    "dev_server": (
        "Re-point the workspace to the correct QA/Prod source before pipeline setup",
        "1. Confirm with the owner whether the dev source is intentional. 2. Update the "
        "Server parameter to QA/Prod. 3. Validate row counts before Deployment Pipeline setup.",
        "Baselines measured against dev SQL are not representative; pipelines need a stable source.",
    ),
}


def analyse_performance(model, table_flags):
    """Aggregate per-table flags into a deduplicated, prioritised issue list."""
    flag_tables = {}
    for tbl, flags in table_flags.items():
        for f in flags:
            flag_tables.setdefault(f, []).append(tbl)

    # incremental refresh: detect absence of RangeStart/RangeEnd anywhere
    has_range = any(
        "RangeStart" in partition_expression(t) for t in user_tables(model)
    )
    if not has_range:
        flag_tables.setdefault("no_incremental", [])

    # duplicate facts
    dups = detect_duplicate_facts(model)
    if dups:
        flag_tables.setdefault("dup_fact", [])
        for a, b in dups:
            flag_tables["dup_fact"].extend([a, b])

    # dev server param
    for name, val in parameters(model):
        if "dev" in val.lower():
            flag_tables.setdefault("dev_server", []).append(name)

    issues = []
    order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    pid = 1
    for key in ISSUE_LIBRARY:
        if key not in flag_tables:
            continue
        title, root_tpl, impact, prio, colour = ISSUE_LIBRARY[key]
        tbls = sorted(set(flag_tables[key]))
        root = root_tpl.format(tables=", ".join(tbls) if tbls else "the model")
        issues.append(
            dict(key=key, title=title, root=root, impact=impact, prio=prio, colour=colour)
        )
    issues.sort(key=lambda i: order.get(i["prio"], 9))
    for i in issues:
        i["id"] = f"P{pid}"
        pid += 1
    return issues


# --------------------------------------------------------------------------- #
#  Rendering: DOCX                                                            #
# --------------------------------------------------------------------------- #
def render_docx(report_name, model, table_flags, issues, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    def shade(cell, colour):
        if not colour:
            return
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), colour)
        tcPr.append(sh)

    def add_table(headers, rows, colour_idx=None):
        """rows: list of tuples; optionally each tuple's last element is a colour."""
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            c.text = h
            shade(c, HEADER)
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
        for row in rows:
            colour = None
            if colour_idx is not None and len(row) > len(headers):
                colour = row[colour_idx]
                row = row[:len(headers)]
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                if colour:
                    shade(cells[i], colour)
        doc.add_paragraph()

    # ---- Heading ----
    doc.add_heading(f"{report_name}", level=1)
    doc.add_paragraph(
        "Source: generated from the Power BI semantic-model JSON. "
        "Amber rows indicate structural inefficiencies; red rows indicate critical blockers."
    )

    ut = sorted(user_tables(model), key=lambda x: x["name"].lower())
    ncols = sum(len(t.get("columns", [])) for t in ut)

    # ---- Table Inventory ----
    doc.add_heading(f"Table Inventory ({len(ut)} Tables, {ncols} Columns)", level=2)
    inv_rows = []
    for t in ut:
        cols = ", ".join(column_label(c) for c in t.get("columns", []))
        typ = classify_table(t, model)
        colour = AMBER if typ in ("Bridge", "Fact", "Lookup") else None
        inv_rows.append((t["name"], typ, cols, colour))
    add_table(["Table", "Type", "Columns"], inv_rows, colour_idx=3)

    # ---- Measure Inventory ----
    doc.add_heading("DAX Measure Inventory", level=2)
    ms = measures(model)
    if ms:
        add_table(["Measure Name", "Table", "Expression"], ms)
    else:
        doc.add_paragraph("No DAX measures defined in the model.")

    # ---- Relationship Map ----
    doc.add_heading("Relationship Map", level=2)
    doc.add_paragraph("Red = many-to-many (performance risk). Amber = structural inefficiency. Green = clean.")
    rel = relationship_rows(model)
    add_table(["Relationship", "Cardinality", "Issue"],
              [(a, b, c, col) for (a, b, c, col) in rel], colour_idx=3)

    # ---- Parameters ----
    doc.add_heading("Parameters", level=2)
    add_table(["Parameter", "Current Value"], parameters(model))

    # ---- Transformation Inventory ----
    doc.add_heading("Transformation Inventory", level=1)
    trows, _ = transformation_rows(model)
    add_table(["Table", "Source", "Logic Type", "Key Transformation Detail"], trows)

    # ---- Performance Analysis ----
    doc.add_heading("Performance Analysis", level=1)
    doc.add_paragraph(
        f"{len(issues)} distinct performance issues identified from the model and "
        "transformation inventory. Red = critical blockers. Amber = significant inefficiencies."
    )
    pa_rows = [
        (i["id"], i["title"], i["root"], i["impact"], i["prio"], i["colour"])
        for i in issues
    ]
    add_table(["#", "Issue", "Root Cause", "Impact", "Priority"], pa_rows, colour_idx=5)

    # ---- Mitigations ----
    doc.add_heading("Performance Mitigations", level=1)
    doc.add_paragraph("Each mitigation is mapped to the issue it resolves, ordered by impact.")
    mit_rows = []
    for i in issues:
        mit, how, why = MITIGATION_LIBRARY[i["key"]]
        mit_rows.append((i["id"], mit, how, why))
    add_table(["Issue", "Mitigation", "How", "Why"], mit_rows)

    # ---- Report metadata (placeholders not in model) ----
    doc.add_heading("Report Metadata (not in model — to be completed)", level=1)
    meta = [
        ("Report name", report_name),
        ("Tables", str(len(ut))),
        ("Columns", str(ncols)),
        ("Measures", str(len(ms))),
        ("Views (30 days)", "<from usage inventory>"),
        ("Users (30 days)", "<from usage inventory>"),
        ("Owner", "<from Power BI Service>"),
        ("Workspace", "<from Power BI Service>"),
    ]
    add_table(["Field", "Value"], meta)

    doc.save(out_path)


# --------------------------------------------------------------------------- #
#  Rendering: Markdown                                                        #
# --------------------------------------------------------------------------- #
def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |",
           "| " + " | ".join("---" for _ in headers) + " |"]
    for r in rows:
        out.append("| " + " | ".join(str(c).replace("|", "\\|") for c in r) + " |")
    return "\n".join(out) + "\n"


def render_md(report_name, model, table_flags, issues, out_path):
    ut = sorted(user_tables(model), key=lambda x: x["name"].lower())
    ncols = sum(len(t.get("columns", [])) for t in ut)
    L = [f"# {report_name}\n",
         "_Generated from the Power BI semantic-model JSON._\n"]

    L.append(f"## Table Inventory ({len(ut)} Tables, {ncols} Columns)\n")
    L.append(md_table(["Table", "Type", "Columns"],
             [(t["name"], classify_table(t, model),
               ", ".join(column_label(c) for c in t.get("columns", []))) for t in ut]))

    L.append("## DAX Measure Inventory\n")
    ms = measures(model)
    L.append(md_table(["Measure", "Table", "Expression"], ms) if ms else "_No measures._\n")

    L.append("## Relationship Map\n")
    L.append(md_table(["Relationship", "Cardinality", "Issue"],
             [(a, b, c) for (a, b, c, _) in relationship_rows(model)]))

    L.append("## Parameters\n")
    L.append(md_table(["Parameter", "Current Value"], parameters(model)))

    L.append("## Transformation Inventory\n")
    trows, _ = transformation_rows(model)
    L.append(md_table(["Table", "Source", "Logic Type", "Key Transformation Detail"], trows))

    L.append("## Performance Analysis\n")
    L.append(md_table(["#", "Issue", "Root Cause", "Impact", "Priority"],
             [(i["id"], i["title"], i["root"], i["impact"], i["prio"]) for i in issues]))

    L.append("## Performance Mitigations\n")
    mit = []
    for i in issues:
        m, how, why = MITIGATION_LIBRARY[i["key"]]
        mit.append((i["id"], m, how, why))
    L.append(md_table(["Issue", "Mitigation", "How", "Why"], mit))

    Path(out_path).write_text("\n".join(L), encoding="utf-8")


# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser(description="Generate a BRD-style report section from a Power BI model JSON.")
    ap.add_argument("json", help="Path to the semantic-model JSON (e.g. database.json)")
    ap.add_argument("-o", "--output", help="Explicit output file path (overrides --output-dir)")
    ap.add_argument("--output-dir", default="reports",
                    help="Folder to write reports into (default: reports). "
                         "Files are named after the report.")
    ap.add_argument("--format", choices=["docx", "md"], default="docx")
    args = ap.parse_args()

    model, name = load_model(args.json)
    report_name = re.sub(r"[ _]+", " ", name).strip()

    _, table_flags = transformation_rows(model)
    issues = analyse_performance(model, table_flags)

    if args.output:
        out = Path(args.output)
    else:
        out = Path(args.output_dir) / f"{safe_filename(report_name)}.{args.format}"
    out.parent.mkdir(parents=True, exist_ok=True)
    out = str(out)

    if args.format == "docx":
        render_docx(report_name, model, table_flags, issues, out)
    else:
        render_md(report_name, model, table_flags, issues, out)

    print(f"✓ Generated {args.format.upper()} report: {out}")
    print(f"  Report:        {report_name}")
    print(f"  Tables:        {len(user_tables(model))} user tables")
    print(f"  Measures:      {len(measures(model))}")
    print(f"  Relationships: {len(relationship_rows(model))}")
    print(f"  Issues found:  {len(issues)}  ({', '.join(i['id']+'='+i['prio'] for i in issues)})")


if __name__ == "__main__":
    main()
